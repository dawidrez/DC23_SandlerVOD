from .collect_data import collect_data

# Generates a document with statistics from the given date range in the given format
def generate_document(start_date, end_date, format):
    data = collect_data(start_date, end_date)
    if format == 'docx':
        return generate_docx(data, start_date, end_date)
    elif format == 'odt':
        return generate_odt(data, start_date, end_date)
    
    raise ValueError('Unsupported format: ' + format)

def generate_docx(data, start_date, end_date):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import matplotlib.pyplot as plt
    import os
    from datetime import datetime

    def add_heading(doc, text, level):
        doc.add_heading(text, level=level)

    def add_paragraph(doc, text, style=None):
        if not isinstance(text, str):
            text = str(text)
        return doc.add_paragraph(text, style=style)

    def add_image(doc, image_path, width):
        doc.add_picture(image_path, width=width)

    def set_cell_border(cell, border_color="000000", border_width="4"):
        """
        Ustawia obramowanie i wyśrodkowanie dla komórki tabeli.

        Args:
            cell: komórka tabeli
            border_color: kolor obramowania w formacie RGB (domyślnie czarny)
            border_width: szerokość obramowania (domyślnie 4)
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Ustawianie obramowania
        tcBorders = OxmlElement('w:tcBorders')
        for border in ['top', 'left', 'bottom', 'right']:
            borderEl = OxmlElement(f'w:{border}')
            borderEl.set(qn('w:val'), 'single')
            borderEl.set(qn('w:sz'), border_width)
            borderEl.set(qn('w:color'), border_color)
            tcBorders.append(borderEl)
        tcPr.append(tcBorders)

        # Ustawianie wyrównania pionowego (wyśrodkowanie)
        tcVertAlign = OxmlElement('w:vAlign')
        tcVertAlign.set(qn('w:val'), 'center')
        tcPr.append(tcVertAlign)

        tcHorizAlign = OxmlElement('w:jc')
        tcHorizAlign.set(qn('w:val'), 'center')
        tcPr.append(tcHorizAlign)

        # Ustawianie wyrównania poziomego (wyśrodkowanie)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_timestamp(doc):
        # Tworzenie akapitu z tekstem "Wygenerowano..." wyrównanym do prawej
        paragraph = doc.add_paragraph()
        run = paragraph.add_run('Wygenerowano: ' + datetime.now().strftime('%d.%m.%Y %H:%M:%S'))
        
        # Ustawienie mniejszej wielkości czcionki
        font = run.font
        font.size = Pt(10)

        # Wyrównanie akapitu do prawej
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def add_table(doc, headers, data):
        table = doc.add_table(rows=1, cols=len(headers))

        # Ustawienie stylu nagłówka tabeli
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell_paragraph = cell.paragraphs[0]
            cell_run = cell_paragraph.add_run(header)
            cell_run.bold = True
            set_cell_border(cell)

        # Dodawanie wierszy danych
        for row_data in data:
            row = table.add_row()
            for i, cell_data in enumerate(row_data):
                cell = row.cells[i]
                cell_paragraph = cell.paragraphs[0]
                cell_paragraph.add_run(str(cell_data))
                set_cell_border(cell)

    def create_pie_chart(labels, sizes, file_name):
        fig, ax = plt.subplots()
        ax.pie(sizes, labels=labels, autopct='%1.1f%%', textprops={'fontsize': 14}, pctdistance=.6, labeldistance=6.6)
        ax.axis('equal')  # Zapewnienie, że wykres jest kołem
        ax.legend(loc="upper right")
        plt.savefig(file_name, format='png', transparent=True)
        plt.close()

    doc = Document()

    data_converted = data['average_duration']
    data_converted = str(data_converted)
    data_converted = data_converted.split(':')
    days = int(data_converted[0].split(' ')[0])
    hours = int(data_converted[0].split(' ')[-1])
    data_converted = str(days) + ' dni, ' + str(hours) + ' godzin i ' + data_converted[1] + ' minut'

    # Styl tekstu
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    # Nagłówek dokumentu
    add_timestamp(doc)
    add_heading(doc, 'Raport miesięczny SandlerVOD', 0)
    add_paragraph(doc, 'Za okres od ' + start_date.strftime('%d.%m.%Y') + ' do ' + end_date.strftime('%d.%m.%Y'))

    # Sekcje i tabele
    sections = [
        ('Liczba klientów subskrybujących każdy pakiet', ["Nazwa pakietu", "Liczba klientów"], data['clients_count_by_package']),
        ('Analiza przychodów według rodzaju pakietu', ["Nazwa pakietu", "Przychód"], data['income_by_package']),
        ('Liczba aktywnych użytkowników', ["Liczba aktywnych użytkowników"], data['active_clients_count']),
        ('Liczba nowych subskrypcji', ["Liczba nowych subskrypcji"], data['new_clients_count']),
        ('Średni czas trwania subskrypcji', ["Średni czas trwania subskrypcji"], data_converted),
        ('Najpopularniejsze filmy w pakietach', ["Tytuł filmu", "Liczba subskrypcji"], data['popular_movies']),
    ]

    labels = []
    sizes = []
    for key, value in data['clients_count_by_package'].items():
        labels.append(key)
        sizes.append(value)
    chart_file0 = f'pie_chart0.png'
    create_pie_chart(labels, sizes, chart_file0)

    labels = []
    sizes = []
    for key, value in data['income_by_package'].items():
        labels.append(key)
        sizes.append(value)
    chart_file1 = f'pie_chart1.png'
    create_pie_chart(labels, sizes, chart_file1)

    labels = []
    sizes = []
    for key, value in data['popular_movies'].items():
        labels.append(key)
        sizes.append(value)
    chart_file5 = f'pie_chart5.png'
    create_pie_chart(labels, sizes, chart_file5)



    for i, (title, headers, table_data) in enumerate(sections):
        add_heading(doc, title, level=1)
        add_paragraph(doc, '')
        if isinstance(table_data, dict):
            table_data_formatted = [[key, str(value)] for key, value in table_data.items()]
            add_table(doc, headers, table_data_formatted)
            add_paragraph(doc, '')
            add_image(doc, 'pie_chart' + str(i) + '.png', width=Inches(6))
            add_paragraph(doc, '')
        else:
            add_paragraph(doc, str(table_data), style='List Bullet')

    # Zapisywanie dokumentu
    doc.save('report.docx')

    # Usuń plik pie_chart.png, jeśli istnieje
    if os.path.exists('pie_chart0.png'):
        os.remove('pie_chart0.png')
    if os.path.exists('pie_chart1.png'):
        os.remove('pie_chart1.png')
    if os.path.exists('pie_chart5.png'):
        os.remove('pie_chart5.png')

    return 'report.docx'


def generate_odt(data, start_date, end_date):
    from odf.opendocument import OpenDocumentText
    from odf.style import Style, TextProperties, ParagraphProperties, TableColumnProperties, TableRowProperties, TableCellProperties, TableProperties
    from odf.text import H, P, Span
    from odf.table import Table, TableRow, TableCell, TableColumn
    from odf.draw import Frame, Image
    import matplotlib.pyplot as plt
    from datetime import datetime
    import os

    def create_style(doc, name, family, textalign, fontsize, fontweight=None, marginleft=None, marginright=None):
        style = Style(name=name, family=family)
        style.addElement(ParagraphProperties(textalign=textalign, marginleft=marginleft, marginright=marginright))
        text_props = {'fontsize': fontsize}
        if fontweight:
            text_props['fontweight'] = fontweight
        style.addElement(TextProperties(attributes=text_props))
        doc.styles.addElement(style)
        return style

    def create_table_styles(doc):
        # Styl tabeli z marginesami
        table_style = Style(name="TableStyle", family="table")
        table_style.addElement(TableColumnProperties(breakbefore="auto"))
        table_style.addElement(TableRowProperties(keeptogether="auto"))
        table_style.addElement(TableProperties(marginleft="0.1in", marginright="0.1in"))
        doc.automaticstyles.addElement(table_style)

        # Styl komórki tabeli
        table_cell_style = Style(name="pourcent", family="table-cell")
        table_cell_style.addElement(TableCellProperties(border="0.74pt solid #000000", padding="0.1in"))
        doc.automaticstyles.addElement(table_cell_style)

        return table_style, table_cell_style

    def add_heading(doc, text, style):
        heading = H(outlinelevel=2, stylename=style)
        heading.addElement(Span(text=text))
        doc.text.addElement(heading)

    def add_result(doc, text):
        result = P(text=text, stylename=result_style)
        doc.text.addElement(result)

    add_footer_center = lambda doc, text: add_heading(doc, text, footer_style_center)
    add_footer_right = lambda doc, text: add_heading(doc, text, footer_style_right)

    def add_image(doc, image_path, frame_width, frame_height, frame_x, frame_y):
        # Dodaj obraz do dokumentu jako zasób binarny
        image_uri = doc.addPicture(image_path)

        # Utwórz ramkę i dodaj obraz
        frame = Frame(width=frame_width, height=frame_height, x=frame_x, y=frame_y)
        image_element = Image(href=image_uri)
        frame.addElement(image_element)
        doc.text.addElement(frame)

    def add_table(doc, headers, data, table_style, table_cell_style):
        table = Table(stylename=table_style)

        tr = TableRow()
        for header in headers:
            tc = TableCell(stylename=table_cell_style)
            p = P(text=header)
            tc.addElement(p)
            tr.addElement(tc)
        table.addElement(tr)

        # Dodawanie danych
        for row in data:
            tr = TableRow()
            for cell in row:
                tc = TableCell(stylename=table_cell_style)
                p = P(text=str(cell))
                tc.addElement(p)
                tr.addElement(tc)
            table.addElement(tr)

        doc.text.addElement(table)

    def add_spacing(doc):
        # Dodaje pusty akapit jako separator
        spacing = P(text="")
        doc.text.addElement(spacing)


    doc = OpenDocumentText()

    data_converted = data['average_duration']
    data_converted = str(data_converted)
    data_converted = data_converted.split(':')
    days = int(data_converted[0].split(' ')[0])
    hours = int(data_converted[0].split(' ')[-1])
    data_converted = str(days) + ' dni, ' + str(hours) + ' godzin i ' + data_converted[1] + ' minut'
    
    # Tworzenie stylów
    title_style = create_style(doc, 'Title', 'paragraph', 'center', '24pt', 'bold')
    heading_style = create_style(doc, 'Heading', 'paragraph', 'center', '18pt', 'bold')
    result_style = create_style(doc, 'Result', 'paragraph', 'center', '14pt', 'bold')
    footer_style_center = create_style(doc, 'FooterCenter', 'paragraph', 'center' , '10pt')
    footer_style_right = create_style(doc, 'FooterRight', 'paragraph', 'right' , '10pt')
    table_style, table_cell_style = create_table_styles(doc)

    # Dodawanie nagłówka dokumentu
    add_footer_right(doc, 'Wygenerowano: ' + datetime.now().strftime('%d.%m.%Y %H:%M:%S'))
    add_spacing(doc)
    add_heading(doc, 'Raport miesięczny SandlerVOD', title_style)
    add_spacing(doc)
    add_footer_center(doc, 'Za okres od ' + start_date.strftime('%d.%m.%Y') + ' do ' + end_date.strftime('%d.%m.%Y'))
    add_spacing(doc)
    add_spacing(doc)
    add_spacing(doc)
    add_spacing(doc)

    # Sekcje i tabele
    sections = [
        ('Liczba klientów subskrybujących każdy pakiet', ["Nazwa pakietu", "Liczba klientów"], data['clients_count_by_package']),
        ('Analiza przychodów według rodzaju pakietu', ["Nazwa pakietu", "Przychód"], data['income_by_package']),
        ('Liczba aktywnych użytkowników', ["Liczba aktywnych użytkowników"], data['active_clients_count']),
        ('Liczba nowych subskrypcji', ["Liczba nowych subskrypcji"], data['new_clients_count']),
        ('Średni czas trwania subskrypcji', ["Średni czas trwania subskrypcji"], data_converted),
        ('Najpopularniejsze filmy w pakietach', ["Tytuł filmu", "Liczba subskrypcji"], data['popular_movies']),
    ]

    labels = []
    sizes = []
    for key, value in data['clients_count_by_package'].items():
        labels.append(key)
        sizes.append(value)

    fig, ax = plt.subplots()
    ax.pie(sizes, labels=labels, autopct='%1.1f%%', textprops={'fontsize': 14}, pctdistance=.6, labeldistance=6.6)
    ax.axis('equal')  # Zapewnienie, że wykres jest kołem
    ax.legend(loc="upper right")
    # Zapisz wykres jako obraz
    plt.savefig('pie_chart0.png', transparent=True)
    plt.close()

    labels = []
    sizes = []
    for key, value in data['income_by_package'].items():
        labels.append(key)
        sizes.append(value)

    fig, ax = plt.subplots()
    ax.pie(sizes, labels=labels, autopct='%1.1f%%', textprops={'fontsize': 14}, pctdistance=.6, labeldistance=6.6)
    ax.axis('equal')  # Zapewnienie, że wykres jest kołem
    ax.legend(loc="upper right")
    # Zapisz wykres jako obraz
    plt.savefig('pie_chart1.png', transparent=True)
    plt.close()

    labels = []
    sizes = []
    for key, value in data['popular_movies'].items():
        labels.append(key)
        sizes.append(value)

    fig, ax = plt.subplots()
    ax.pie(sizes, labels=labels, autopct='%1.1f%%', textprops={'fontsize': 14},pctdistance=.6, labeldistance=6.6)
    ax.axis('equal')  # Zapewnienie, że wykres jest kołem
    ax.legend(loc="upper right")
    # Zapisz wykres jako obraz
    plt.savefig('pie_chart5.png', transparent=True)
    plt.close()

    
    for i, (title, headers, table_data) in enumerate(sections):
        add_spacing(doc)
        add_heading(doc, title, heading_style)
        add_spacing(doc)
        table_data_formatted = []
        if isinstance(table_data, dict):
            # Dla słowników, iteruj po ich elementach
            for key, value in table_data.items():
                table_data_formatted.append([key, value])
            add_table(doc, headers, table_data_formatted, table_style, table_cell_style)
            add_spacing(doc)
            add_image(doc, 'pie_chart' + str(i) + '.png', '6in', '4in', '0in', '0in')
        else:
            if table_data == 0:
                add_result(doc, 'Brak danych')
            else:
                add_result(doc, table_data)
            add_spacing(doc)

    # Zapisywanie dokumentu
    doc.save('report.odt')

    # Usuń plik pie_chart.png, jeśli istnieje
    if os.path.exists('pie_chart0.png'):
        os.remove('pie_chart0.png')
    if os.path.exists('pie_chart1.png'):
        os.remove('pie_chart1.png')
    if os.path.exists('pie_chart5.png'):
        os.remove('pie_chart5.png')

    return 'report.odt'


